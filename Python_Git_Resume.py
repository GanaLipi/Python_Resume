#importing python docx to create, read and Write doc files
import docx

#created a doc object
doc = docx.Document()

#In 'section' Headers and footers are linked.
sec =doc.sections[0]

#Creating a header and footer object
header = sec.header
footer = sec.footer

#adding title or writing a name into the headers and footers
title   = header.paragraphs[0]
title_1 = footer.paragraphs[0]

title.add_run ('Enter Your Name').bold = True
title.add_run ('\t \tGitHub:').bold = True
title.add_run (' https://www.github.com/GanaLipi\n') 
title.add_run ('Enter Your Email ID')
title.add_run ('\nEnter Your Phone Number')

title_1.text = '\t \tThis Resume is created using Python'

#Heading is added with its name
doc.add_heading ('Education').bold = True

#adding a table to doc and created a table object
table = doc.add_table (rows =4, cols = 4)

#accessing each cells using rows and colms
cell_00 = table.cell(0,0)
cell_00.text = 'Qualification'
cell_10 = table.cell(1,0)
cell_10.text = '         BE'
cell_20 = table.cell(2,0)
cell_20.text = '    Diploma'
cell_30 = table.cell(3,0)
cell_30.text = '       SSLC'

cell_01 = table.cell(0,1)
cell_01.text = '     Year'
cell_11 = table.cell(1,1)
cell_11.text = '2010-2014'
cell_21 = table.cell(2,1)
cell_21.text = '2006-2010'
cell_31 = table.cell(3,1)
cell_31.text = '2005-2006'

cell_02 = table.cell(0,2)
cell_02.text = '     University'
cell_12 = table.cell(1,2)
cell_12.text = '  VTU, Belagavi'
cell_22 = table.cell(2,2)
cell_22.text = ' DTE, Bengaluru'
cell_32 = table.cell(3,2)
cell_32.text = 'KSSEB, Bengaluru'

cell_03 = table.cell(0,3)
cell_03.text = 'Percentage'
cell_13 = table.cell(1,3)
cell_13.text = '     60.0%'
cell_23 = table.cell(2,3)
cell_23.text = '     70.0%'
cell_33 = table.cell(3,3)
cell_33.text = '     75.0%'



doc.add_heading ('Projects').bold = True

table_1 = doc.add_table (rows = 3, cols = 2)

#accesssing each cells of a colomns
heading_cells = table_1.rows[0].cells
heading_cells[0].text = '\tProject Title'
heading_cells[1].text = '\tProject Description'

heading_cells = table_1.rows[1].cells
heading_cells[0].text = '\tProject Title'
heading_cells[1].text = '\tProject Description'
heading_cells = table_1.rows[2].cells
heading_cells[0].text = '\tProject Title'
heading_cells[1].text = '\tProject Description'

#IF you have experiene you can add anywhere you need
#doc.add_heading ('Experience').bold = True
#doc.add_paragraph ('Company')
#doc.add_paragraph ('Designation')
#doc.add_paragraph ('Experience')

doc.add_heading ('Skills').bold = True
#adding a paragraph to the document
doc.add_paragraph ('Languages')
doc.add_paragraph ('Communication Protocols')
doc.add_paragraph ('Controllers')
doc.add_paragraph ('Web Skills')

#adding a page break
#doc.add_page_break()

doc.add_heading ('Certification').bold = True
doc.add_paragraph ('1. Enter The Certifications you took for Example')
doc.add_paragraph ('2. PG Certificate on Full Stack from xxx institution')

doc.add_heading ('Achievements and Awards').bold = True
doc.add_paragraph ('1. Enter the recent Awards you got for Example')
doc.add_paragraph ('2. Won 1st Prize in PaperPresentation on Topic Virtual Reality')

doc.add_heading ('Extra information')
doc.add_paragraph ('1. Add the Extra Curriculam Activities for Example')
doc.add_paragraph ('2. I was managed few Events in my engineering days')

#Saving a file, can save in different directory by providing Path 
doc.save('Python_GIT_RESUME.docx')
